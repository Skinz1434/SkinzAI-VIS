"""
Document Processing Pipeline
Handles OCR, classification, and text extraction from various formats
"""

import asyncio
from typing import List, Dict, Any, Optional, BinaryIO
from pathlib import Path
import hashlib
import io
from datetime import datetime
import re
import json

import PyPDF2
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image
import numpy as np
import cv2
from docx import Document as DocxDocument
import fitz  # PyMuPDF for better PDF handling
from pdfplumber import PDF

from app.core.logging import logger
from app.models.documents import DocumentType, ProcessingStatus

class DocumentProcessor:
    """Main document processing pipeline"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.tiff']
        self.ocr_config = '--oem 3 --psm 6'  # OCR Engine Mode 3, Page Segmentation Mode 6
        
    async def process_document(self, file) -> Dict:
        """
        Process a single document through the pipeline
        """
        logger.info(f"Processing document: {file.filename}")
        
        # Generate document ID
        doc_id = self._generate_doc_id(file.filename)
        
        # Read file content
        content = await file.read()
        
        # Determine file type
        file_extension = Path(file.filename).suffix.lower()
        
        # Process based on type
        if file_extension == '.pdf':
            result = await self._process_pdf(content, doc_id)
        elif file_extension == '.docx':
            result = await self._process_docx(content, doc_id)
        elif file_extension == '.txt':
            result = await self._process_text(content, doc_id)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff']:
            result = await self._process_image(content, doc_id)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Classify document type
        result["type"] = await self._classify_document(result["text"])
        
        # Extract metadata
        result["metadata"] = await self._extract_metadata(result["text"])
        
        # Add processing info
        result.update({
            "id": doc_id,
            "filename": file.filename,
            "format": file_extension,
            "processed_at": datetime.utcnow().isoformat(),
            "status": ProcessingStatus.COMPLETED
        })
        
        logger.info(f"Document processed successfully: {doc_id}")
        return result
    
    def _generate_doc_id(self, filename: str) -> str:
        """Generate unique document ID"""
        timestamp = datetime.utcnow().isoformat()
        hash_input = f"{filename}{timestamp}".encode()
        return hashlib.sha256(hash_input).hexdigest()[:16]
    
    async def _process_pdf(self, content: bytes, doc_id: str) -> Dict:
        """Process PDF document"""
        result = {
            "pages": [],
            "text": "",
            "has_images": False,
            "needs_ocr": False
        }
        
        try:
            # Try text extraction first
            with fitz.open(stream=content, filetype="pdf") as pdf:
                for page_num, page in enumerate(pdf, 1):
                    # Extract text
                    page_text = page.get_text()
                    
                    # Check if page has meaningful text
                    if len(page_text.strip()) < 50:
                        # Page might be scanned, needs OCR
                        result["needs_ocr"] = True
                        
                        # Convert page to image for OCR
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scaling for better OCR
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        
                        # Preprocess image for OCR
                        img = await self._preprocess_image_for_ocr(img)
                        
                        # Perform OCR
                        page_text = pytesseract.image_to_string(img, config=self.ocr_config)
                    
                    # Store page data
                    result["pages"].append({
                        "page_number": page_num,
                        "text": page_text,
                        "has_images": len(page.get_images()) > 0
                    })
                    
                    result["text"] += f"\n[Page {page_num}]\n{page_text}\n"
                
                # Check for images
                if any(p["has_images"] for p in result["pages"]):
                    result["has_images"] = True
        
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            # Fallback to pure OCR
            result = await self._process_pdf_with_ocr(content, doc_id)
        
        return result
    
    async def _process_pdf_with_ocr(self, content: bytes, doc_id: str) -> Dict:
        """Process PDF using OCR for all pages"""
        result = {
            "pages": [],
            "text": "",
            "has_images": True,
            "needs_ocr": True
        }
        
        # Convert PDF to images
        images = convert_from_bytes(content, dpi=300)
        
        for page_num, img in enumerate(images, 1):
            # Preprocess for OCR
            img = await self._preprocess_image_for_ocr(img)
            
            # Perform OCR with error handling
            try:
                page_text = pytesseract.image_to_string(img, config=self.ocr_config)
                
                # Clean OCR output
                page_text = await self._clean_ocr_text(page_text)
                
                result["pages"].append({
                    "page_number": page_num,
                    "text": page_text,
                    "ocr_confidence": await self._calculate_ocr_confidence(page_text)
                })
                
                result["text"] += f"\n[Page {page_num}]\n{page_text}\n"
                
            except Exception as e:
                logger.error(f"OCR failed for page {page_num}: {e}")
                result["pages"].append({
                    "page_number": page_num,
                    "text": "[OCR Failed]",
                    "error": str(e)
                })
        
        return result
    
    async def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Preprocess image to improve OCR accuracy"""
        # Convert PIL to OpenCV
        img_array = np.array(image)
        
        # Convert to grayscale
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Apply thresholding to get black and white image
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Denoise
        denoised = cv2.fastNlMeansDenoising(thresh, h=30)
        
        # Deskew
        angle = self._get_skew_angle(denoised)
        if abs(angle) > 0.5:
            denoised = self._rotate_image(denoised, angle)
        
        # Convert back to PIL
        return Image.fromarray(denoised)
    
    def _get_skew_angle(self, image: np.ndarray) -> float:
        """Detect skew angle of scanned document"""
        edges = cv2.Canny(image, 50, 200)
        lines = cv2.HoughLinesP(edges, 1, np.pi/180, 100, minLineLength=100, maxLineGap=10)
        
        if lines is None:
            return 0.0
        
        angles = []
        for line in lines:
            x1, y1, x2, y2 = line[0]
            angle = np.degrees(np.arctan2(y2 - y1, x2 - x1))
            angles.append(angle)
        
        if angles:
            median_angle = np.median(angles)
            return median_angle if abs(median_angle) < 10 else 0.0
        return 0.0
    
    def _rotate_image(self, image: np.ndarray, angle: float) -> np.ndarray:
        """Rotate image to correct skew"""
        h, w = image.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        return rotated
    
    async def _clean_ocr_text(self, text: str) -> str:
        """Clean and normalize OCR output"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR errors
        replacements = {
            'l)': 'b',  # Common OCR mistake
            '()': 'O',
            '|': 'I',
            '0': 'O',  # In certain contexts
        }
        
        for old, new in replacements.items():
            # Context-aware replacement would go here
            pass
        
        # Remove special characters that are likely OCR artifacts
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\/\@\#\$\%\&\*\+\=]', '', text)
        
        return text.strip()
    
    async def _calculate_ocr_confidence(self, text: str) -> float:
        """Calculate confidence score for OCR output"""
        # Simple heuristic based on text quality
        confidence = 1.0
        
        # Check for excessive special characters (likely OCR errors)
        special_char_ratio = len(re.findall(r'[^a-zA-Z0-9\s\.\,]', text)) / max(len(text), 1)
        if special_char_ratio > 0.2:
            confidence -= 0.3
        
        # Check for reasonable word length
        words = text.split()
        if words:
            avg_word_length = sum(len(w) for w in words) / len(words)
            if avg_word_length < 2 or avg_word_length > 15:
                confidence -= 0.2
        
        # Check for dictionary words (would need dictionary integration)
        # For now, just check if text has recognizable patterns
        if not re.search(r'\b[A-Za-z]{3,}\b', text):
            confidence -= 0.3
        
        return max(confidence, 0.1)
    
    async def _process_docx(self, content: bytes, doc_id: str) -> Dict:
        """Process Word document"""
        result = {
            "pages": [],
            "text": "",
            "has_images": False,
            "tables": []
        }
        
        doc = DocxDocument(io.BytesIO(content))
        
        # Extract text from paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        result["text"] = "\n".join(full_text)
        
        # Extract tables
        for table_num, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append({
                "table_number": table_num,
                "data": table_data
            })
        
        # Check for images
        if doc.inline_shapes:
            result["has_images"] = True
        
        # Simulate pages (Word doesn't have explicit pages)
        # Split text into chunks of ~3000 characters per "page"
        page_size = 3000
        text_chunks = [result["text"][i:i+page_size] 
                      for i in range(0, len(result["text"]), page_size)]
        
        for page_num, chunk in enumerate(text_chunks, 1):
            result["pages"].append({
                "page_number": page_num,
                "text": chunk
            })
        
        return result
    
    async def _process_text(self, content: bytes, doc_id: str) -> Dict:
        """Process plain text document"""
        text = content.decode('utf-8', errors='ignore')
        
        result = {
            "pages": [{
                "page_number": 1,
                "text": text
            }],
            "text": text,
            "has_images": False
        }
        
        return result
    
    async def _process_image(self, content: bytes, doc_id: str) -> Dict:
        """Process image file"""
        img = Image.open(io.BytesIO(content))
        
        # Preprocess for OCR
        img = await self._preprocess_image_for_ocr(img)
        
        # Perform OCR
        text = pytesseract.image_to_string(img, config=self.ocr_config)
        text = await self._clean_ocr_text(text)
        
        result = {
            "pages": [{
                "page_number": 1,
                "text": text,
                "ocr_confidence": await self._calculate_ocr_confidence(text)
            }],
            "text": text,
            "has_images": True,
            "needs_ocr": True
        }
        
        return result
    
    async def _classify_document(self, text: str) -> str:
        """Classify document type based on content"""
        text_lower = text.lower()
        
        # Check for specific document types
        if "dd form 214" in text_lower or "dd214" in text_lower:
            return DocumentType.DD214
        elif "va form 21-526" in text_lower:
            return DocumentType.VA_FORM_526EZ
        elif "disability benefits questionnaire" in text_lower or "dbq" in text_lower:
            return DocumentType.DBQ
        elif "medical opinion" in text_lower and "nexus" in text_lower:
            return DocumentType.NEXUS_LETTER
        elif "buddy statement" in text_lower or "lay statement" in text_lower:
            return DocumentType.BUDDY_STATEMENT
        elif any(term in text_lower for term in ["diagnosis", "treatment", "medical record"]):
            return DocumentType.MEDICAL_RECORD
        elif "service treatment record" in text_lower or "str" in text_lower:
            return DocumentType.SERVICE_TREATMENT_RECORD
        elif "rating decision" in text_lower:
            return DocumentType.RATING_DECISION
        elif "c&p exam" in text_lower or "compensation and pension" in text_lower:
            return DocumentType.CP_EXAM
        else:
            return DocumentType.OTHER
    
    async def _extract_metadata(self, text: str) -> Dict:
        """Extract metadata from document text"""
        metadata = {}
        
        # Extract dates
        date_pattern = r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2})\b'
        dates = re.findall(date_pattern, text)
        if dates:
            metadata["dates"] = dates[:10]  # Limit to first 10 dates
        
        # Extract SSN (masked)
        ssn_pattern = r'\b\d{3}-\d{2}-\d{4}\b'
        ssns = re.findall(ssn_pattern, text)
        if ssns:
            metadata["ssn_detected"] = True
            metadata["ssn_masked"] = [f"XXX-XX-{ssn[-4:]}" for ssn in ssns]
        
        # Extract names (using simple pattern, would use NER in production)
        name_pattern = r'\b([A-Z][a-z]+ [A-Z][a-z]+)\b'
        names = re.findall(name_pattern, text)
        if names:
            metadata["possible_names"] = list(set(names[:5]))  # Unique, first 5
        
        # Extract medical terms
        medical_terms = self._extract_medical_terms(text)
        if medical_terms:
            metadata["medical_terms"] = medical_terms
        
        # Extract military units
        unit_pattern = r'\b\d{1,3}(?:st|nd|rd|th)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b'
        units = re.findall(unit_pattern, text)
        if units:
            metadata["military_units"] = list(set(units))
        
        return metadata
    
    def _extract_medical_terms(self, text: str) -> List[str]:
        """Extract medical terminology from text"""
        # Common medical terms relevant to VA claims
        medical_keywords = [
            "ptsd", "tbi", "traumatic brain injury", "depression", "anxiety",
            "hearing loss", "tinnitus", "back pain", "knee pain", "sleep apnea",
            "diabetes", "hypertension", "ischemic heart disease", "peripheral neuropathy",
            "radiculopathy", "degenerative disc disease", "arthritis", "migraines",
            "gerd", "ibs", "chronic fatigue", "fibromyalgia", "sciatica"
        ]
        
        found_terms = []
        text_lower = text.lower()
        
        for term in medical_keywords:
            if term in text_lower:
                found_terms.append(term)
        
        # Also extract ICD codes
        icd_pattern = r'\b[A-Z]\d{2}(?:\.\d{1,2})?\b'
        icd_codes = re.findall(icd_pattern, text)
        found_terms.extend(icd_codes)
        
        return list(set(found_terms))
    
    async def batch_process(self, files: List) -> List[Dict]:
        """Process multiple documents in parallel"""
        tasks = [self.process_document(file) for file in files]
        results = await asyncio.gather(*tasks)
        return results